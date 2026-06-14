# -*- coding: utf-8 -*-
"""
產生「2026 慈悲科技創新競賽」Aeterlux 計畫書 .docx
格式遵循主辦規定：A4 直向、內文12級、標題14級、中文標楷體、英文 Times New Roman、
行距1.5倍、邊界上下2.54cm 左右3.18cm。

內容由 content.py 提供（SECTIONS / FIGURES），本檔只負責排版與嵌圖。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

import content  # SECTIONS, FIGURES, TITLE, TEAM

HERE = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(HERE, "_proposal_imgs")
OUT = os.path.join(HERE, "Aeterlux_計畫書_第二組.docx")

KAI = "標楷體"
TNR = "Times New Roman"
CONTENT_WIDTH_CM = 21.0 - 3.18 * 2  # A4寬 - 左右邊界 ≈ 14.64cm


def set_cjk_font(run, font=KAI, size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = TNR  # 西文
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font)  # 中文用標楷體
    rfonts.set(qn("w:ascii"), TNR)
    rfonts.set(qn("w:hAnsi"), TNR)
    if color:
        run.font.color.rgb = color


def set_line_spacing(p, mult=1.5, space_before=0, space_after=6):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = mult
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_line_spacing(p, 1.5, space_before=8, space_after=4)
    run = p.add_run(text)
    set_cjk_font(run, size=14, bold=True)
    return p


def add_body(doc, text, indent=True):
    """一段內文；text 可能含換行，逐段加。"""
    for block in text.split("\n"):
        block = block.strip()
        if not block:
            continue
        p = doc.add_paragraph()
        set_line_spacing(p, 1.5)
        if indent:
            p.paragraph_format.first_line_indent = Pt(24)  # 首行縮排兩字
        # 子標題（以 (一)/（二）/1. 等開頭或以「：」結束的短行）粗體處理：簡化—整段同字型
        run = p.add_run(block)
        set_cjk_font(run, size=12)
    return


def add_subheading(doc, text):
    p = doc.add_paragraph()
    set_line_spacing(p, 1.5, space_before=4, space_after=2)
    run = p.add_run(text)
    set_cjk_font(run, size=12, bold=True)
    return p


def add_figure(doc, filename, caption, width_ratio=0.72):
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path):
        # 缺圖時放提示
        p = doc.add_paragraph()
        set_line_spacing(p, 1.5)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"【缺圖：{filename}】")
        set_cjk_font(run, size=12, color=RGBColor(0xC0, 0x00, 0x00))
        return
    # 依原圖比例計算寬度
    target_w = CONTENT_WIDTH_CM * width_ratio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(p, 1.0, space_before=4, space_after=2)
    run = p.add_run()
    run.add_picture(path, width=Cm(target_w))
    # 圖說
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(cap, 1.0, space_before=0, space_after=8)
    crun = cap.add_run(caption)
    set_cjk_font(crun, size=10.5)


def build():
    doc = Document()
    # 邊界
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)

    # 預設樣式字型
    style = doc.styles["Normal"]
    style.font.size = Pt(12)
    style.font.name = TNR
    style.element.rPr.rFonts.set(qn("w:eastAsia"), KAI)

    # 封面標題（封面不計入10頁，但放在首段作為標題塊）
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(title_p, 1.5, space_before=0, space_after=2)
    trun = title_p.add_run(content.TITLE)
    set_cjk_font(trun, size=16, bold=True)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_line_spacing(sub_p, 1.5, space_before=0, space_after=10)
    srun = sub_p.add_run(content.SUBTITLE)
    set_cjk_font(srun, size=12)

    # 逐段
    for sec_def in content.SECTIONS:
        add_heading(doc, sec_def["heading"])
        for blk in sec_def["blocks"]:
            if blk["type"] == "body":
                add_body(doc, blk["text"], indent=blk.get("indent", True))
            elif blk["type"] == "subheading":
                add_subheading(doc, blk["text"])
            elif blk["type"] == "figure":
                add_figure(doc, blk["file"], blk["caption"], blk.get("width", 0.72))
            elif blk["type"] == "note":
                # 建議補充圖片備註：灰底框感（用斜體+顏色）
                for line in blk["text"].split("\n"):
                    if not line.strip():
                        continue
                    p = doc.add_paragraph()
                    set_line_spacing(p, 1.5, space_after=2)
                    run = p.add_run(line.strip())
                    set_cjk_font(run, size=11, color=RGBColor(0x55, 0x55, 0x55))
                    run.italic = True
            elif blk["type"] == "reference":
                # 引用條目：第一行「[n] 粗體出處名稱」，網址各自一行、縮排小灰字。
                # blk: {"label": "[1] 國發會…", "urls": ["https://…", ...]}
                head = doc.add_paragraph()
                set_line_spacing(head, 1.3, space_before=4, space_after=1)
                hrun = head.add_run(blk["label"])
                set_cjk_font(hrun, size=11, bold=True)
                for url in blk.get("urls", []):
                    up = doc.add_paragraph()
                    set_line_spacing(up, 1.2, space_before=0, space_after=2)
                    up.paragraph_format.left_indent = Pt(18)
                    urun = up.add_run(url)
                    set_cjk_font(urun, size=9.5, color=RGBColor(0x55, 0x60, 0x70))

    doc.save(OUT)
    print("WROTE", OUT)


if __name__ == "__main__":
    build()
